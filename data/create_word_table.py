#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
生成可编辑的 Word (.docx) 表格文件
运行前请确保已安装依赖：pip install pandas python-docx
保存本文件为 create_word_table.py 后运行：python create_word_table.py
"""

import io
import base64
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt

# CSV 文本（根据你提供的表格）
csv_text = """断面名称,项目类型,酸碱度 PH,溶解氧 DO (mg/L),生物需氧量 BOD5 (mg/L),化学需氧量 COD (mg/L),六价铬 Cr+6 (mg/L),氨氮 NH3-N (mg/L),高锰酸盐指数 (mg/L),色度,悬浮物 SS (mg/L),浊度,其它,正磷酸盐 (mg/L)
复兴大桥点1,浓度/(mg/L),6.12,8.9,8.38,17.6,0.061,0.46,3.6981,1,35,9.38,0.090,0.01175
复兴大桥点1,超标倍数,未超标,未超标,1.095,未超标,0.22,未超标,未超标,未超标,/,/,未超标,未超标
复兴大桥点2,浓度/(mg/L),7.20,8.3,3.1825,35.2,0.073,待确认,4.4485,1.5,12,15.38,0.073,0.01318
复兴大桥点2,超标倍数,未超标,未超标,未超標,0.76,0.46,未超标,未超标,未超标,/,/,未超标,未超标
标准值,,6-9,>=5,<=4,<=20,<=0.05,<=1,<=6,<=15,无,无,<=10,<=0.2
"""

def create_docx_from_df(df: pd.DataFrame, out_path: Path):
    doc = Document()
    doc.add_heading("水质表", level=2)

    # 创建表格：行 = header + data rows，列 = dataframe 列数
    rows = len(df) + 1
    cols = len(df.columns)
    table = doc.add_table(rows=rows, cols=cols)
    table.style = 'Table Grid'

    # 写入表头（加粗）
    hdr_cells = table.rows[0].cells
    for j, col in enumerate(df.columns):
        p = hdr_cells[j].paragraphs[0]
        run = p.add_run(str(col))
        run.bold = True
        # 设置表头字体大小（可选）
        run.font.size = Pt(10)

    # 写入数据行
    for i, row in enumerate(df.itertuples(index=False), start=1):
        for j, value in enumerate(row):
            cell = table.rows[i].cells[j]
            cell.text = "" if pd.isna(value) else str(value)

    # 保存
    doc.save(out_path)
    return out_path


def main():
    df = pd.read_csv(io.StringIO(csv_text), dtype=str)  # 以字符串方式读取，保留原样
    out_docx = Path("水质表.docx")
    create_docx_from_df(df, out_docx)
    print(f"已生成 Word 文件: {out_docx.resolve()}")

    # 同时生成 base64 文件
    b = out_docx.read_bytes()
    b64 = base64.b64encode(b).decode("utf-8")
    out_b64 = Path("水质表.docx.b64")
    out_b64.write_text(b64, encoding="utf-8")
    print(f"已生成 base64 文件: {out_b64.resolve()}")
    print("如果你希望我把 base64 内容直接粘在聊天里以便你复制，请回复'是'。")


if __name__ == "__main__":
    main()
